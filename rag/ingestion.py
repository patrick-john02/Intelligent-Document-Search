from langchain_ollama import OllamaEmbeddings
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_core.documents import Document
from pptx import Presentation
from fastapi import File, UploadFile, HTTPException, status
from typing import Dict, Any, AsyncGenerator
from langchain_ollama import OllamaLLM
from dataclasses import dataclass, field
from openpyxl import load_workbook
from PIL import Image
from core.dependencies import Deps, deps
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from core.dependencies import SessionLocal
from api.models.document import(
    DocumentVersion, 
    DocumentModel,
)


import pandas as pd
import uuid
import numpy as np
import asyncio
import io
import anydoc

@dataclass
class FileChunk:
    document_version_id:int
    file_name: str
    chunk_id: str
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)


#Extraction of text we need to identify the extension name
#split the file name and .pdf for example
#i will be using the anydoc library for text extraction, rather than doing if else
async def extract_text(
        file_bytes:bytes,
)->str:

    
    markdown = anydoc.to_markdown_bytes(file_bytes)
    return markdown

#     if extension == 'pdf':
#         pdf_reader = PdfReader(io.BytesIO(file_bytes))
#         text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])

#         if pdf_reader.metadata:
#             metadata = {key.strip('/'): value for key, value in pdf_reader.metadata.items()}


#     elif extension == "docx":
#         word_reader = DocxDocument(io.BytesIO(file_bytes))

#         text = "\n".join([paragraph.text for paragraph in word_reader.paragraphs])

#         core_props = word_reader.core_properties
#         metadata = {
#             "author": core_props.author,
#             "title": core_props.title,
#             "subject": core_props.subject,
#             "created": core_props.created,
#             "modified": core_props.modified
#         }

#         metadata = {key: value for key, value in metadata.items() if value is not None}


#     elif extension == "xlsx":
#         excel_reader = io.BytesIO(file_bytes)
#         dataframe = pd.read_excel(excel_reader)

#         text = dataframe.to_string(index=False)

#         excel_reader.seek(0)
#         workbook = load_workbook(excel_reader, read_only=True)
#         core_props = workbook.properties

#         metadata = {
#             "creator": core_props.creator or "Unknown",
#             "title": core_props.title or "Unknown",
#             "created": core_props.created or "Unknown",
#             "modified": core_props.modified or "Unknown"
#         }

    # if extension in ["jpg", "jpeg", "png"]:
    #     image_reader = Image.open(io.BytesIO(file_bytes))

    #     metadata = {
    #         "format": image_reader.format,
    #         "width": image_reader.width,
    #         "height": image_reader.height,
    #         "mode": image_reader.mode,
    #     }

    # else:
    #     raise ValueError(f"Unsupport file extension {extension}")

    # return text, metadata
    



def chunk_text(
        text:str,
        chunk_size: int = 1000,
        overlap: int = 200 
)->list[str]:
    
    if not text:
        return []

    splitter = RecursiveCharacterTextSplitter(
        chunk_size = chunk_size, 
        chunk_overlap = overlap,
        separators=["\n\n", "\n", " ", ""]

    )
    return splitter.split_text(text)

async def insert_file_chunk(
        sem: asyncio.Semaphore,
        deps:Deps,
        chunk: FileChunk,
)->None:
    async with sem:
        unique_string =  f"{chunk.document_version_id}_chunk_{chunk.chunk_id}"
        point_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, unique_string))
        
        try:
            # embedding_response = await deps.embedding_client.aembed_documents(
            #     [chunk.embedding_content()]
            # )
            
            # vector = embedding_response[0]
            
            document = Document(
                page_content=chunk.content,
                metadata={
                **chunk.metadata,
                "document_version_id": chunk.document_version_id,
                "file_name": chunk.file_name,
                "chunk_id": chunk.chunk_id
                },
        )
            
            await deps.vector_store.aadd_documents(
                documents=[document],
                ids=[point_id],
            )


        except Exception:
            raise

async def process_uploaded_file(document_version_id: int, file_name: str, file_path:str):
    
    try:

        print(f"Extracting Text from {document_version_id}")

        # with open(file_path, "rb") as f:
        #     file_bytes = f.read()
        
        async with SessionLocal() as db_session:
            query = (
                select(DocumentVersion)
                .where(DocumentVersion.id == document_version_id)
                .options(selectinload(DocumentVersion.document))
            )
            result = await db_session.execute(query)
            version = result.scalar_one_or_none()
        
        if not version or not version.document:
            raise ValueError(f"DocumentVersion {document_version_id} or parent Document not found")
        
        document_id = version.document_id
        version_number = version.version_number
        clearance_level = (
            version.document.clearance_level.value
            if hasattr(version.document.clearance_level, "value")
            else str(version.document.clearance_level)
        )
        
        with open(file_path, "rb") as f:
            file_bytes = f.read()
        
        

        # raw_text, metadata = await extract_text(file_name, file_bytes)
        markdown = await extract_text(file_bytes)
        text_chunk = chunk_text(markdown)


        file_chunks = [
            FileChunk(
                document_version_id=document_version_id,
                file_name=file_name, 
                chunk_id=str(i), 
                content=chunk
            )

            for i, chunk in enumerate(text_chunk)

        ]

        sem = asyncio.Semaphore(5)
        task = [insert_file_chunk(sem, deps, chunk) for chunk in file_chunks]
        await asyncio.gather(*task)


        async with SessionLocal() as db_session:
            version = await db_session.get(DocumentVersion, document_version_id)
            if version:
                version.status="indexed"
                await db_session.commit()
                
    except Exception as e:
        print(f"Ingestion Failed for version {document_version_id}:{e}")
        async with SessionLocal() as db_session:
            version = await db_session.get(DocumentVersion, document_version_id)
            if version:
                version.status = "failed"
                await db_session.commit()
                
        


